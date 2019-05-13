# -*- coding: utf-8 -*-

# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: https://doc.scrapy.org/en/latest/topics/item-pipeline.html

import os.path
import sys

sys.path.append('..')


from docx import Document

from data.file_types import FileTypes


class DocxPipeline():

    def __init__(self):
        # make the book title an argument
        self.document = Document()

    def process_item(self, item, spider):
        """Write the chapter to the a Document, which was created from the constructor"""

        # if the file type is not doxc skip this pipeline
        if spider.file_type != FileTypes.DOXC:
            return item

        # add linebreak to the lines and join them to a string
        list_of_lines = item['chapter_content']
        list_of_lines = map(lambda x: x + '\n', list_of_lines)
        chapter_content = ''.join(list_of_lines)

        # Write chapter title
        para = self.document.add_heading(item['chapter_title'])
        # Make sure that the chapter always starts at the top of the next page
        para.paragraph_format.page_break_before = True
        # Write chapter content
        self.document.add_paragraph(chapter_content)
        # End page, start next chapter on the next page
        # self.document.add_page_break()

        return item

    def close_spider(self, spider):
        path = spider.book_name.strip() + '.docx'
        path = os.path.join('Output', 'Docx', path)
        self.document.save(path)
